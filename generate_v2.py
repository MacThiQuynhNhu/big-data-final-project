#!/usr/bin/env python3
"""Generate BAO_CAO_BIGDATA_v2.docx — phong cách luận văn, dài, đầy đủ."""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image as _PILImage
import os, re

DST = 'BAO_CAO_BIGDATA_v2.docx'

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(13)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.3

rPr = style.element.get_or_add_rPr()
rFonts = rPr.makeelement(qn('w:rFonts'), {})
rFonts.set(qn('w:ascii'), 'Times New Roman')
rFonts.set(qn('w:hAnsi'), 'Times New Roman')
rPr.insert(0, rFonts)

for level, size, color in [(1, 18, '2F5496'), (2, 15, '2F5496'), (3, 13, '333333')]:
    hs = doc.styles[f'Heading {level}']
    hf = hs.font
    hf.name = 'Times New Roman'
    hf.size = Pt(size)
    hf.bold = True
    hf.color.rgb = RGBColor.from_string(color)
    hs.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    hs.paragraph_format.space_after = Pt(6)

# ── Helpers ──
def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    return h

def para(text, bold_parts=None):
    """Add paragraph. Use **text** for bold inline."""
    p = doc.add_paragraph()
    parts = text.split('**')
    for i, part in enumerate(parts):
        if not part:
            continue
        run = p.add_run(part)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(13)
        if i % 2 == 1:
            run.bold = True
    return p

def bullet(text):
    p = doc.add_paragraph(text, style='List Bullet')
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(13)
    return p

def numbered(text):
    p = doc.add_paragraph(text, style='List Number')
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(13)
    return p

def table(headers, rows):
    """Create a Word table with headers and data rows."""
    ncols = len(headers)
    nrows = len(rows) + 1
    tbl = doc.add_table(rows=nrows, cols=ncols)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        cell = tbl.cell(0, j)
        cell.text = ''
        r = cell.paragraphs[0].add_run(h)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(11)
        r.bold = True
        shading = cell._element.get_or_add_tcPr().makeelement(qn('w:shd'), {})
        shading.set(qn('w:fill'), 'D9D9D9')
        shading.set(qn('w:val'), 'clear')
        cell._element.get_or_add_tcPr().append(shading)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = tbl.cell(i + 1, j)
            cell.text = ''
            parts = str(val).split('**')
            for pi, part in enumerate(parts):
                if not part: continue
                r = cell.paragraphs[0].add_run(part)
                r.font.name = 'Times New Roman'
                r.font.size = Pt(11)
                if pi % 2 == 1: r.bold = True
    doc.add_paragraph()
    return tbl

def code(text):
    for line in text.strip().split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        run.font.name = 'Courier New'
        run.font.size = Pt(10)
        shd = p.paragraph_format.element.get_or_add_pPr().makeelement(qn('w:shd'), {})
        shd.set(qn('w:fill'), 'F5F5F5')
        shd.set(qn('w:val'), 'clear')
        p.paragraph_format.element.get_or_add_pPr().append(shd)

def img_placeholder(text):
    p = doc.add_paragraph()
    run = p.add_run(f'[ {text} ]')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    run.italic = True
    run.font.color.rgb = RGBColor.from_string('888888')

def add_image(fname, caption=None, cap_in=6.3, folder='report_images'):
    """Chèn ảnh nếu có (canh giữa, có caption). Thiếu file -> giữ caption placeholder (an toàn)."""
    path = fname if os.path.exists(fname) else os.path.join(folder, fname)
    if os.path.exists(path):
        try:
            w, _ = _PILImage.open(path).size
            width_in = min(cap_in, w / 96.0)
        except Exception:
            width_in = cap_in
        doc.add_picture(path, width=Inches(width_in))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            c = doc.add_paragraph()
            c.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = c.add_run('Hình: ' + caption)
            r.font.name = 'Times New Roman'; r.font.size = Pt(11); r.italic = True
            r.font.color.rgb = RGBColor.from_string('555555')
    else:
        img_placeholder('Chụp ảnh: ' + (caption or fname))

def quote(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.5)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.italic = True
    run.font.color.rgb = RGBColor.from_string('555555')

def hr():
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pBdr.makeelement(qn('w:bottom'), {})
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)

def ascii_art(text):
    code(text)

def _field(p, instr, placeholder=''):
    r = p.add_run()
    f1 = OxmlElement('w:fldChar'); f1.set(qn('w:fldCharType'), 'begin')
    it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = instr
    f2 = OxmlElement('w:fldChar'); f2.set(qn('w:fldCharType'), 'separate')
    t = OxmlElement('w:t'); t.text = placeholder
    f3 = OxmlElement('w:fldChar'); f3.set(qn('w:fldCharType'), 'end')
    for e in (f1, it, f2, t, f3):
        r._r.append(e)
    return r

def add_toc():
    """Mục lục tự động — trong Word nhấn Ctrl+A rồi F9 để cập nhật số trang."""
    _field(doc.add_paragraph(), 'TOC \\o "1-3" \\h \\z \\u',
           'Nhấn Ctrl+A rồi F9 để cập nhật mục lục.')

def add_page_numbers():
    """Footer: 'Trang X' canh giữa."""
    p = doc.sections[0].footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('Trang ')
    _field(p, 'PAGE')
    for run in p.runs:
        run.font.name = 'Times New Roman'; run.font.size = Pt(10)

def page_break():
    from docx.enum.text import WD_BREAK
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

# ══════════════════════════════════════════
# CONTENT
# ══════════════════════════════════════════

# ── Cover ──
heading('BÁO CÁO CUỐI KỲ MÔN BIG DATA', 1)
heading('TỔNG HỢP BÁO CÁO KINH DOANH ĐA NGUỒN', 2)
doc.add_paragraph()
para('**Học viên:** [Họ tên] — **Mã HV:** [MSSV]')
para('**Giảng viên:** [Tên giảng viên]')
para('**Ngày:** 06/2026')
hr()

# ── Mục lục ──
page_break()
heading('MỤC LỤC', 1)
add_toc()
page_break()

# ── PART 1 ──
heading('PHẦN 1 — THIẾT KẾ HỆ THỐNG', 1)

# 1.1
heading('1.1. Bài toán kinh doanh đa nguồn', 2)

para('Trong bối cảnh chuyển đổi số đang diễn ra mạnh mẽ trên toàn cầu, các doanh nghiệp bán lẻ không còn giới hạn ở một kênh bán hàng duy nhất. Thay vào đó, họ mở rộng sang nhiều kênh cùng lúc — từ cửa hàng vật lý truyền thống cho đến website thương mại điện tử và ứng dụng di động. Mỗi kênh bán hàng này lại được vận hành bởi một hệ thống công nghệ thông tin riêng biệt, được thiết kế cho một mục đích cụ thể.')

para('Hậu quả trực tiếp của mô hình này là dữ liệu kinh doanh bị phân mảnh nghiêm trọng. Một giao dịch bán hàng có thể được ghi nhận tại hệ thống POS ở quầy thu ngân, đồng thời được phản ánh trong hệ thống ERP để hạch toán tài chính, trong khi thông tin khách hàng lại nằm ở một hệ thống CRM hoàn toàn độc lập. Mỗi hệ thống sử dụng định dạng dữ liệu, tần suất cập nhật và ngữ nghĩa riêng, khiến cho việc tổng hợp báo cáo kinh doanh toàn diện trở thành một thách thức lớn.')

para('Cụ thể, một doanh nghiệp bán lẻ điển hình thường vận hành các hệ thống sau:')

table(
    ['Hệ thống', 'Kênh', 'Vai trò'],
    [
        ['**POS** (Point of Sale)', 'Offline', 'Ghi nhận hóa đơn tại quầy thu ngân, hoạt động theo thời gian thực'],
        ['**ERP** (Enterprise Resource Planning)', 'Offline + Online', 'Hệ thống tài chính tổng hợp, ghi nhận toàn bộ giao dịch từ mọi kênh, có cột kenh phân biệt online/offline'],
        ['**E-commerce**', 'Online', 'Quản lý đơn hàng từ website và ứng dụng di động, lưu thông tin thiết bị và phương thức thanh toán'],
        ['**Kho/WMS** (Warehouse Management)', '—', 'Quản lý tồn kho, theo dõi từng hoạt động nhập và xuất hàng hóa'],
        ['**CRM** (Customer Relationship Management)', '—', 'Quản lý thông tin khách hàng, phân khúc và lịch sử tương tác'],
    ]
)

para('**Vấn đề** đặt ra là: làm thế nào để thu thập dữ liệu từ tất cả các hệ thống này một cách tự động, làm sạch và chuẩn hóa về một định dạng chung, sau đó phân tích để đưa ra các báo cáo kinh doanh có ý nghĩa? Đây chính là bài toán mà hệ thống được trình bày trong báo cáo này hướng đến giải quyết. Cụ thể, hệ thống cần đáp ứng các yêu cầu sau:')

bullet('Thu thập dữ liệu từ tất cả các nguồn theo thời gian thực, không can thiệp thủ công')
bullet('Làm sạch và chuẩn hóa dữ liệu về một schema thống nhất, bất kể định dạng đầu vào')
bullet('Phân tích doanh thu, chi phí, lợi nhuận và tồn kho theo nhiều chiều thời gian: ngày, tuần, tháng, quý và năm')
bullet('Phát hiện và cảnh báo các giao dịch bất thường theo thời gian thực')
bullet('Trực quan hóa kết quả phân tích trên dashboard, có thể truy cập mọi lúc')

# 1.2
heading('1.2. Kiến trúc hệ thống', 2)

para('Để giải quyết bài toán đã nêu, hệ thống được thiết kế với kiến trúc hai lớp xử lý song song — Batch và Streaming — cùng chia sẻ một tầng dữ liệu chung. Đây là biến thể của kiến trúc Lambda đã được chứng minh hiệu quả trong nhiều hệ thống Big Data trên thế giới. Ý tưởng cốt lõi là: dữ liệu được đưa vào một hàng đợi trung tâm (Kafka), từ đó phân phối đồng thời đến hai nhánh xử lý — một nhánh cho phân tích định kỳ chính xác cao (batch) và một nhánh cho giám sát thời gian thực (streaming).')

ascii_art("""┌──────────────────────────────────────────────────────────────────┐
│                    5 DATA SOURCES                                │
│   POS (API)  │  ERP (DB)  │  ECOM (DB)  │  KHO (DB)  │  CRM (API) │
└───────────────────────┬──────────────────────────────────────────┘
                        │
                 ┌──────▼──────┐
                 │  Apache NiFi │  ← Thu thập + Làm sạch + Định tuyến
                 └──────┬──────┘
                        │
                 ┌──────▼──────┐
                 │ Apache Kafka │  ← Hàng đợi tin nhắn (2 topics)
                 └──┬──────┬───┘
                    │      │
        ┌───────────▼──┐ ┌─▼────────────────┐
        │ BATCH LAYER  │ │  STREAMING LAYER │
        │ Spark SQL    │ │ Spark Streaming  │
        │ (Cron 15')   │ │ (Real-time)      │
        │ HDFS /lake   │ │ Kafka → rt_*     │
        │   ↓          │ │   ↓              │
        │ Hive bao_cao │ │ PostgreSQL       │
        └──────┬───────┘ └──────┬───────────┘
               │                │
        ┌──────▼────────────────▼──────────┐
        │        SERVING LAYER             │
        │   PostgreSQL (marts) + Grafana   │
        │   http://192.168.79.131:3000     │
        └──────────────────────────────────┘""")

para('**Batch Layer** chịu trách nhiệm xử lý dữ liệu theo định kỳ. Spark đọc các file JSON từ HDFS /lake — nơi NiFi đã đổ dữ liệu từ Kafka xuống — sau đó ghi vào Hive dưới dạng bảng có cấu trúc, phân vùng theo nguồn và tháng. Từ đây, các truy vấn SQL và mô hình học máy được thực thi để tạo ra các báo cáo tổng hợp (bc_*) và dữ liệu tích lũy theo thời gian (agg_*). Toàn bộ quá trình này được lập lịch chạy tự động mỗi 15 phút thông qua Cron.')

para('**Streaming Layer** đọc trực tiếp từ Kafka để tính toán các chỉ số theo thời gian thực. Không giống như batch layer cần đợi đến chu kỳ tiếp theo, streaming layer xử lý từng micro-batch dữ liệu chỉ trong vài giây, cho phép phát hiện ngay các giao dịch bất thường và cập nhật dashboard tức thì. Kết quả được ghi vào các bảng rt_thongke và rt_canhbao trong PostgreSQL.')

para('**Serving Layer** là tầng phục vụ dữ liệu cho người dùng cuối. Grafana kết nối trực tiếp đến PostgreSQL, truy vấn đồng thời cả dữ liệu batch (bc_*, agg_*) và streaming (rt_*) để hiển thị trên cùng một dashboard. Người dùng có thể xem báo cáo tổng quan theo năm, quý, hoặc kéo timeline để phân tích chi tiết đến từng ngày — tất cả đều tự động cập nhật mà không cần thao tác thủ công.')

add_image('kien_truc_chuan.png', 'Sơ đồ kiến trúc hệ thống (Lambda — Batch + Streaming)')

# 1.3
heading('1.3. Cụm Hadoop và công nghệ sử dụng', 2)

para('Toàn bộ hệ thống được triển khai trên một cụm Hadoop gồm 3 máy ảo VMware, mô phỏng môi trường phân tán thực tế. Mặc dù tài nguyên phần cứng có hạn (tổng 12GB RAM), cụm vẫn vận hành đầy đủ các thành phần của một hệ thống Big Data tiêu chuẩn.')

heading('Cấu hình cụm', 3)

table(
    ['Node', 'IP', 'RAM', 'Vai trò'],
    [
        ['**master**', '10.0.2.195', '6 GB', 'NameNode, ResourceManager, NiFi, Kafka, PostgreSQL, Grafana'],
        ['**slave01**', '10.0.2.196', '3 GB', 'DataNode, NodeManager'],
        ['**slave02**', '10.0.2.197', '3 GB', 'DataNode, NodeManager'],
    ]
)

table(
    ['Cấu hình HDFS', 'Giá trị'],
    [
        ['Replication', '2'],
        ['Block size', '128 MB'],
        ['Tổng dung lượng', '~38 GB (configured)'],
    ]
)

para('Việc phân bổ 6GB cho master và 3GB cho mỗi slave là cân nhắc kỹ lưỡng dựa trên vai trò của từng node. Master đảm nhiệm nhiều service cùng lúc — NameNode cho HDFS, ResourceManager cho YARN, cùng với Kafka, NiFi, PostgreSQL và Grafana — do đó cần nhiều RAM hơn. Hai slave chủ yếu chạy DataNode và NodeManager, đảm nhiệm việc lưu trữ và tính toán phân tán. Với cấu hình này, cụm có thể xử lý đồng thời batch processing trên YARN với 2 executor (mỗi executor 640MB) mà không làm gián đoạn các service khác.')

heading('Công nghệ sử dụng', 3)

table(
    ['Thành phần', 'Công nghệ', 'Phiên bản'],
    [
        ['Thu thập dữ liệu', 'Apache NiFi', '1.28.1'],
        ['Hàng đợi tin nhắn', 'Apache Kafka (KRaft)', '3.7.1'],
        ['Lưu trữ phân tán', 'HDFS (Hadoop)', '3.3.0'],
        ['Kho dữ liệu', 'Apache Hive (Metastore MySQL)', '3.1.1'],
        ['Xử lý theo lô', 'Apache Spark SQL', '3.1.1'],
        ['Học máy', 'Apache Spark MLlib', '3.1.1'],
        ['Xử lý luồng', 'Spark Streaming', '3.1.1'],
        ['Cơ sở dữ liệu nguồn', 'PostgreSQL', '14.x'],
        ['Trực quan hóa', 'Grafana', '10.x'],
        ['Điều phối', 'Bash + Cron (mỗi 15 phút)', '—'],
        ['Mô phỏng dữ liệu', 'Python (source_feeder.py)', '3.x'],
    ]
)

# 1.4
heading('1.4. Mô phỏng dữ liệu đầu vào', 2)

para('Dữ liệu được xây dựng dựa trên bộ Superstore — một tập dữ liệu bán lẻ phổ biến trên Kaggle, chứa thông tin về đơn hàng, sản phẩm, khách hàng và khu vực trong khoảng thời gian 4 năm. Đây là cơ sở ban đầu để làm giàu và phát sinh dữ liệu cho 5 hệ thống nguồn. Script Python source_feeder.py đảm nhiệm việc biến đổi tập dữ liệu tĩnh này thành một dòng dữ liệu liên tục, mô phỏng đúng cách mỗi hệ thống doanh nghiệp thật vận hành.')

heading('Cơ chế mô phỏng', 3)

table(
    ['Tham số', 'Giá trị'],
    [
        ['Tốc độ', '**3 phút thật = 1 ngày mô phỏng**'],
        ['Ngày bắt đầu', '2026-07-01'],
        ['Dữ liệu gốc', 'Dataset Superstore (Kaggle)'],
        ['Phương thức sinh', 'POS: API HTTP POST | ERP/Ecom/Kho: INSERT PostgreSQL | CRM: Ghi file JSON'],
    ]
)

heading('Cách sinh từ dữ liệu gốc', 3)

table(
    ['Nguồn', 'Cách sinh từ dữ liệu gốc', 'Đặc điểm'],
    [
        ['**POS**', 'Tách mỗi đơn hàng Superstore thành hóa đơn nhiều dòng, gán store_id — gửi qua HTTP POST tới NiFi ListenHTTP (port 9998)', 'Hóa đơn nhiều dòng, mỗi lần 1-5 hóa đơn'],
        ['**ERP**', 'Sao chép toàn bộ giao dịch, bổ sung cost (giá vốn) và cột kenh (online/offline) — INSERT vào bảng sales', 'Giao dịch tài chính có revenue, cost, kenh'],
        ['**E-commerce**', 'Chọn subset đơn hàng, thêm device, payment_method — INSERT vào bảng ecommerce_orders', 'Đơn online có device, payment_method'],
        ['**Kho/WMS**', 'Khởi tạo tồn kho ban đầu, phát sinh nhập kho khi tồn chạm ngưỡng, xuất kho theo bán hàng — INSERT vào bảng kho_chuyendong', 'Nhập kho (qty>0) khi tồn dưới ngưỡng, xuất kho (qty<0) theo bán'],
        ['**CRM**', 'Tạo danh sách khách hàng từ customer_id, gán phân khúc ngẫu nhiên; khi có khách mới — gửi qua API CRM. E-commerce chọn ngẫu nhiên customer_id từ danh sách này khi tạo đơn', 'Khách hàng mới dần, liên kết với đơn online'],
    ]
)

heading('Kết quả', 3)

table(
    ['Chỉ số', 'Giá trị'],
    [
        ['Tổng giao dịch ERP (đã phân tích)', '**198,840**'],
        ['Phạm vi phân tích', '2026-07-01 → 2027-04-04 (253 ngày, ~9 tháng)'],
        ['Giao dịch offline (POS)', '126,616'],
        ['Giao dịch online (Ecom)', '72,224'],
        ['Số khách hàng định danh (CRM)', '53'],
    ]
)

quote('Feeder được thiết kế để chạy liên tục, ngày mô phỏng tiến dần, nối tiếp dữ liệu cũ để tạo thành timeline liền mạch. Có thể dừng và khởi động lại bất kỳ lúc nào mà không gây trùng lặp txn_id. Số liệu phân tích trong toàn bộ báo cáo dựa trên 198,840 giao dịch từ nguồn ERP (chốt ở các kỳ đã đóng tới 2027-04-04).')

para('Thông qua cơ chế trên, từ một tập dữ liệu tĩnh ban đầu, hệ thống đã tạo ra một dòng dữ liệu liên tục mô phỏng đúng đặc điểm của từng hệ thống doanh nghiệp: POS đẩy hóa đơn theo thời gian thực, ERP ghi nhận tài chính cho cả hai kênh, E-commerce tích lũy đơn hàng trực tuyến, Kho theo dõi biến động tồn kho, và CRM cập nhật thông tin khách hàng mới.')

# 1.5
heading('1.5. Các nguồn dữ liệu', 2)

para('Hệ thống thu thập dữ liệu từ 5 nguồn độc lập, mỗi nguồn có định dạng và phương thức truy cập riêng. Bảng dưới đây tổng hợp thông tin về từng nguồn, bao gồm cả số bản ghi đã được ingest vào Hive.')

table(
    ['#', 'Nguồn', 'Kênh', 'Định dạng', 'NiFi Processor', 'Số bản ghi'],
    [
        ['1', '**POS**', 'Offline', 'HTTP API (JSON)', 'ListenHTTP :9998', '126,616'],
        ['2', '**ERP**', 'Cả 2', 'PostgreSQL sales', 'QueryDatabaseTable', '198,840'],
        ['3', '**E-commerce**', 'Online', 'PostgreSQL ecommerce_orders', 'QueryDatabaseTable', '72,224'],
        ['4', '**Kho/WMS**', '—', 'PostgreSQL kho_chuyendong', 'QueryDatabaseTable', '9 (tồn kho)'],
        ['5', '**CRM**', '—', 'HTTP API (JSON)', 'InvokeHTTP :8000', '53'],
    ]
)

quote('Mối quan hệ giữa các nguồn: ERP là hệ thống tài chính, ghi nhận lại toàn bộ giao dịch từ POS (offline) và E-commerce (online). Do đó ERP (198,840) bằng tổng của POS (126,616) và Ecommerce (72,224). POS và Ecom cung cấp thông tin vận hành bổ sung (store_id, device, payment_method...), còn ERP là nguồn chính thức cho mọi phân tích tài chính (có revenue, cost, kenh). CRM liên kết với E-commerce qua customer_id; khách hàng mua tại quầy không có định danh nên không liên kết được.')

heading('1. POS (Point of Sale)', 3)
para('Hệ thống POS ghi nhận giao dịch tại quầy thu ngân. Mỗi hóa đơn bao gồm nhiều dòng sản phẩm và được gửi đến NiFi qua HTTP POST request. Đây là nguồn dữ liệu real-time, phản ánh trực tiếp hoạt động bán hàng tại cửa hàng. POS cung cấp thông tin store_id — mã cửa hàng — giúp phân tích doanh thu theo vị trí địa lý. Dữ liệu mẫu:')
code('{"invoice_id": "POS-00001", "store_id": "California",\n "txn_date": "2026-07-01", "total": 900.0,\n "items": [{"product_id": "TEC-CO-10004722", "qty": 2, "price": 450.0}]}')
add_image('pos_sample.png', 'POS đẩy hóa đơn qua HTTP POST tới NiFi ListenHTTP')

heading('2. ERP (Tài chính)', 3)
para('ERP là nguồn dữ liệu quan trọng nhất trong toàn bộ hệ thống. Đây là hệ thống tài chính, ghi nhận lại toàn bộ giao dịch từ tất cả các kênh bán hàng. Điểm đặc biệt của ERP là nó chứa cột kenh để phân biệt giao dịch online và offline, cùng với các cột revenue (doanh thu) và cost (giá vốn) — hai chỉ số cốt lõi cho mọi phân tích tài chính. Do ERP đã bao phủ toàn bộ giao dịch, spark_incremental.py chỉ sử dụng dữ liệu từ nguồn ERP khi tính toán doanh thu và lợi nhuận, tránh việc đếm trùng với POS và Ecom.')
add_image('erp_sales.png', 'Bảng sales (ERP) — có revenue, cost, kenh online/offline')

heading('3. E-commerce', 3)
para('Hệ thống thương mại điện tử quản lý các đơn hàng phát sinh từ website và ứng dụng di động. Mỗi đơn hàng được lưu vào bảng ecommerce_orders với order_id tăng dần, cho phép NiFi QueryDatabaseTable phát hiện bản ghi mới một cách hiệu quả. Ngoài thông tin giao dịch cơ bản, E-commerce còn cung cấp device (thiết bị mua hàng) và payment_method (phương thức thanh toán) — những thông tin không có ở kênh offline.')
add_image('ecommerce_orders.png', 'Bảng ecommerce_orders — có device, payment_method')

heading('4. Kho/WMS', 3)
para('Hệ thống quản lý kho theo dõi từng hoạt động nhập và xuất hàng hóa. Thay vì chỉ lưu số tồn hiện tại như một con số tĩnh, hệ thống ghi nhận từng sự kiện chuyển động kho: nhập kho (qty > 0) và xuất kho (qty < 0). Tồn kho hiện tại được tính bằng tổng của tất cả các chuyển động: stock_qty = SUM(qty). Cách tiếp cận này cho phép truy vết lịch sử tồn kho theo thời gian và tính giá trị tồn (stock_qty × unit_cost) để liên kết với báo cáo tài chính.')
add_image('kho_chuyendong.png', 'Bảng kho_chuyendong — sự kiện nhập/xuất kho')

heading('5. CRM', 3)
para('Hệ thống CRM quản lý thông tin khách hàng, bao gồm mã khách hàng (customer_id), tên, phân khúc (Consumer, Corporate, Home Office) và khu vực. Khi có khách hàng mới đăng ký, thông tin được gửi qua API và NiFi dùng InvokeHTTP định kỳ đồng bộ dữ liệu. E-commerce khi tạo đơn hàng sẽ liên kết với một customer_id từ danh sách CRM, cho phép phân tích doanh thu theo phân khúc khách hàng. Cần lưu ý rằng khách hàng mua tại quầy (POS offline) không có định danh, do đó chỉ có thể phân tích hành vi khách hàng đối với kênh online.')
add_image('crm_api.png', 'API CRM trả về danh sách khách hàng + phân khúc')

heading('Mối liên kết giữa các nguồn', 3)
para('Bảng dưới đây tóm tắt mối quan hệ giữa 5 nguồn dữ liệu, cho thấy chúng liên kết với nhau qua những cột nào:')

table(
    ['', 'POS', 'E-commerce', 'ERP', 'Kho', 'CRM'],
    [
        ['**ERP**', 'Cùng txn_id', 'Cùng txn_id', '—', 'Qua product_id (giá vốn)', 'Qua Ecom (customer_id)'],
        ['**Kho**', 'Qua product_id', 'Qua product_id', 'Qua product_id', '—', '—'],
        ['**CRM**', 'Không', '**Qua customer_id**', 'Một phần (online)', 'Không', '—'],
    ]
)

heading('Đa dạng phương thức thu thập', 3)
para('Một điểm đáng chú ý trong thiết kế hệ thống là việc khai thác nhiều phương thức ingest khác nhau của Apache NiFi. Mỗi nguồn dữ liệu được thu thập bằng processor phù hợp nhất với đặc điểm của nó:')

table(
    ['Processor', 'Dùng cho', 'Lý do lựa chọn'],
    [
        ['ListenHTTP', 'POS', 'Nhận POST hóa đơn real-time do máy bán hàng chi nhánh chủ động đẩy lên (port 9998)'],
        ['InvokeHTTP', 'CRM', 'Chủ động gọi API CRM định kỳ để đồng bộ danh sách khách hàng — sự kiện rời rạc, API nhẹ'],
        ['QueryDatabaseTable', 'ERP, E-commerce, Kho', 'Dữ liệu dạng bảng tích lũy — phát hiện bản ghi mới qua ID tăng dần, chỉ ingest phần chưa xử lý'],
        ['ConsumeKafka', 'Tất cả nguồn sau làm sạch', 'Hàng đợi tin nhắn phân tán — nhận dữ liệu đã qua làm sạch, ghi xuống HDFS Data Lake'],
    ]
)

quote('Hệ thống này được xây dựng ở mức tối thiểu cho bài tập lớn môn Big Data — 5 nguồn, 9 sản phẩm, 4 khu vực, 12 cửa hàng, ~199,000 giao dịch ERP. Trong thực tế, một hệ thống doanh nghiệp thật sẽ phức tạp hơn nhiều: hàng trăm cửa hàng, hàng nghìn sản phẩm, hàng triệu giao dịch mỗi ngày, data lake dung lượng terabyte. Tuy nhiên, kiến trúc và nguyên lý xử lý — từ thu thập, làm sạch, lưu trữ phân tầng, phân tích batch kết hợp streaming, đến trực quan hóa — là hoàn toàn giống với hệ thống thật.')

# ── PART 2 ──
heading('PHẦN 2 — PIPELINE XỬ LÝ DỮ LIỆU', 1)

# 2.1
heading('2.1. Luồng xử lý dữ liệu', 2)

para('Sau khi dữ liệu được sinh ra từ 5 nguồn, nó đi qua một chuỗi xử lý gồm nhiều bước trước khi đến tay người dùng dưới dạng báo cáo và dashboard. Sơ đồ dưới đây mô tả toàn bộ hành trình của dữ liệu trong hệ thống:')

ascii_art("""Feeder (Python) → 5 nguồn PostgreSQL + file JSON
     │
     ▼
Apache NiFi (5 processors)
     │  ListenHTTP (POS) ┐
     │  QueryDatabaseTable (ERP) ├─→ Kafka topic sales-report-clean
     │  QueryDatabaseTable (Ecom) ┘
     │  QueryDatabaseTable (Kho) ──→ Kafka topic inventory-events
     │  InvokeHTTP (CRM) ──────────→ HDFS /clean/crm/
     │
     ▼
Apache Kafka (2 topics)
     │
     ▼
ConsumeKafka → HDFS /lake/
     ├── /lake/transactions/     ← JSON Lines (mỗi dòng 1 giao dịch)
     ├── /lake/inventory/        ← Dữ liệu nhập xuất kho
     └── /clean/crm/              ← Khách hàng JSON
     │
     ▼
Spark Batch (Cron mỗi 15 phút)
     ├── [1] spark_to_hive.py       → sales_report + inventory + dim + snapshot
     ├── [2] spark_incremental.py   → agg_ngay/tuan/thang/quy/nam
     ├── [3] spark_report_hive.py   → bc_doanhthu_cuahang, bc_loinhuan_vung...
     ├── [4] spark_analysis.py      → MLlib: dự báo, phân cụm, kế hoạch nhập
     └── [5] spark_marts_to_pg.py   → Đẩy 22 bảng sang PostgreSQL
     │
     ▼
PostgreSQL (marts) → Grafana (trực quan hóa)""")

# 2.2
heading('2.2. Thu thập dữ liệu — NiFi → Kafka → HDFS', 2)

para('Bước đầu tiên trong pipeline là thu thập dữ liệu từ 5 nguồn. Apache NiFi đảm nhiệm vai trò này với 5 processor chuyên biệt, mỗi processor được cấu hình để kết nối đến một nguồn dữ liệu cụ thể.')

heading('NiFi Flow', 3)
para('Flow được định nghĩa trong file thu-thap-da-nguon.xml và triển khai trên NiFi instance chạy tại master node. Các processor và luồng dữ liệu của chúng được mô tả trong bảng dưới đây:')

table(
    ['Processor', 'Đọc từ', 'Ghi đến'],
    [
        ['ListenHTTP (port 9998)', 'POS API POST request', 'Kafka sales-report-clean'],
        ['QueryDatabaseTable (ERP)', 'PostgreSQL sales', 'Kafka sales-report-clean'],
        ['QueryDatabaseTable (Ecom)', 'PostgreSQL ecommerce_orders', 'Kafka sales-report-clean'],
        ['QueryDatabaseTable (Kho)', 'PostgreSQL kho_chuyendong', 'Kafka inventory-events'],
        ['InvokeHTTP (port 8000)', 'CRM API', 'HDFS /clean/crm/'],
    ]
)

add_image('nifi_flow.png', 'NiFi flow — 5 processor thu thập đa nguồn')

heading('Kafka Topics', 3)
para('Sau khi được NiFi thu thập, dữ liệu được đưa vào Apache Kafka — một hệ thống message queue phân tán. Kafka hoạt động như một lớp đệm trung tâm, tách biệt quá trình thu thập dữ liệu khỏi quá trình xử lý. Điều này mang lại hai lợi ích quan trọng: thứ nhất, nếu hệ thống xử lý phía sau gặp sự cố, dữ liệu vẫn được lưu đệm trong Kafka và không bị mất; thứ hai, nhiều consumer có thể cùng đọc một topic, cho phép cả batch layer và streaming layer cùng tiêu thụ dữ liệu từ một nguồn duy nhất.')

table(
    ['Topic', 'Nội dung', 'Partition'],
    [
        ['sales-report-clean', 'Giao dịch từ POS + ERP + Ecom', '1'],
        ['inventory-events', 'Hoạt động nhập xuất kho', '1'],
    ]
)

heading('Từ Kafka xuống HDFS', 3)
para('NiFi processor ConsumeKafka đọc message từ 2 topic trên và ghi xuống HDFS dưới dạng JSON Lines — mỗi dòng trong file là một bản ghi độc lập. Cách tổ chức này cho phép Spark đọc dữ liệu một cách hiệu quả sau này.')

bullet('ConsumeKafka (topic sales-report-clean) → /lake/transactions/ — mỗi file JSON chứa nhiều dòng giao dịch')
bullet('ConsumeKafka (topic inventory-events) → /lake/inventory/ — mỗi file JSON chứa các sự kiện nhập/xuất kho')

para('Đây là điểm hợp nhất dữ liệu quan trọng nhất trong toàn bộ kiến trúc: 3 nguồn bán hàng khác nhau (POS, ERP, Ecom) cùng được đổ vào một Kafka topic, sau đó cùng được ghi vào một thư mục HDFS, và cuối cùng cùng được nạp vào một bảng Hive duy nhất. Chính thiết kế này giúp cho việc phân tích đa nguồn trở nên khả thi mà không cần xử lý từng nguồn riêng lẻ.')

heading('HDFS Data Lake — Phân tầng', 3)
code('/lake/\n├── transactions/              ← JSON Lines từ Kafka (đã xử lý → archive)\n├── transactions_archive/      ← 17,483 files đã ingest (~111 MB)\n├── inventory/                 ← Hoạt động nhập xuất kho\n└── /clean/crm/                ← Khách hàng')
add_image('hdfs_lake.png', 'HDFS Data Lake — duyệt /lake trên NameNode UI')

# 2.3
heading('2.3. Xử lý theo lô — Spark → Hive', 2)

para('Đây là bước quan trọng nhất trong toàn bộ batch layer. Spark đọc dữ liệu thô từ HDFS, chuẩn hóa schema từ 3 nguồn khác nhau (POS, ERP, Ecom), và lưu vào Hive dưới dạng bảng có cấu trúc với phân vùng theo nguồn và tháng.')

heading('Cơ chế Incremental Ingest', 3)
para('Một trong những cải tiến quan trọng của hệ thống là cơ chế incremental ingest. Thay vì đọc lại toàn bộ file trong /lake/transactions/ mỗi lần batch chạy — điều sẽ khiến thời gian xử lý tăng tuyến tính theo lượng dữ liệu — hệ thống chỉ đọc những file mới xuất hiện kể từ lần batch trước. Sau khi xử lý xong, các file này được chuyển sang thư mục archive. Nhờ vậy, dù hệ thống chạy trong bao lâu, mỗi lần batch vẫn chỉ cần xử lý một lượng nhỏ dữ liệu mới phát sinh.')

code('''# Chỉ đọc file hiện có (file cũ đã archive từ lần trước)
tx = spark.read.json("hdfs://master:9000/lake/transactions")
# APPEND thay vì OVERWRITE
has_sales = "sales_report" in [t.name for t in spark.catalog.listTables("bao_cao")]
sales.write.mode("append" if has_sales else "overwrite")
    .partitionBy("source", "thang").saveAsTable("bao_cao.sales_report")
# Archive file đã xử lý
subprocess.check_call(["hdfs", "dfs", "-mv",
    "hdfs://master:9000/lake/transactions/*",
    "hdfs://master:9000/lake/transactions_archive/"])''')

heading('Bảng sales_report', 3)
para('Bảng sales_report là kết quả của quá trình chuẩn hóa dữ liệu từ 3 nguồn POS, ERP và Ecom về một schema chung. Đây là bảng trung tâm của toàn bộ hệ thống phân tích — mọi báo cáo, tổng hợp và mô hình học máy sau này đều truy vấn từ bảng này.')

table(
    ['Nguồn (source)', 'Vai trò', 'Số dòng'],
    [
        ['POS', 'Vận hành offline (có store_id)', '126,616'],
        ['E-commerce', 'Vận hành online (có device, payment_method)', '72,224'],
        ['ERP', 'Tài chính — bản sao của POS + Ecom (có revenue, cost, kenh)', '198,840'],
        ['**Tổng số dòng**', 'Mỗi giao dịch xuất hiện ở cả nguồn vận hành + ERP', '**397,680**'],
    ]
)

quote('sales_report chứa ~398k dòng nhưng tương đương 198,840 giao dịch thực tế. Mỗi giao dịch được lưu 2 lần: một lần ở nguồn vận hành (POS hoặc Ecom) để có thông tin chi tiết như store_id và device, một lần ở ERP để có số liệu tài chính như revenue và cost. spark_incremental.py khi tính toán doanh thu chỉ sử dụng WHERE source=\'erp\' để tránh đếm trùng.')

heading('Schema bảng sales_report', 3)
table(
    ['Cột', 'Kiểu', 'Ý nghĩa'],
    [
        ['txn_id', 'string', 'Mã giao dịch (duy nhất trong mỗi nguồn)'],
        ['store_id', 'string', 'Mã cửa hàng (chỉ có ở POS)'],
        ['product_id', 'string', 'Mã sản phẩm'],
        ['region', 'string', 'Khu vực (East/South/Central/West)'],
        ['qty', 'int', 'Số lượng bán'],
        ['revenue', 'double', 'Doanh thu'],
        ['cost', 'double', 'Giá vốn (COGS)'],
        ['kenh', 'string', 'Kênh bán: online hoặc offline'],
        ['customer_id', 'string', 'Mã khách hàng (từ Ecom + CRM)'],
        ['device', 'string', 'Thiết bị mua (chỉ có ở Ecom)'],
        ['payment_method', 'string', 'Phương thức thanh toán (chỉ có ở Ecom)'],
        ['txn_date', 'date', 'Ngày giao dịch'],
        ['source', 'string', 'Nguồn gốc: pos, erp, ecommerce'],
        ['thang', 'int', 'Tháng (1-12), dùng làm partition key'],
    ]
)

para('Về cột kenh: ERP đã có sẵn cột này để phân biệt online và offline cho cả hai kênh. Với POS và Ecom, Spark tự suy ra: POS luôn là offline, Ecom luôn là online. Khi cùng một giao dịch xuất hiện ở cả ERP và POS/Ecom, hệ thống ưu tiên giá trị kenh từ ERP.')

quote('Ghi chú: các ảnh truy vấn Hive trong báo cáo được chụp bằng hàm rút gọn hq() { spark-sql --master local[1] -e "$1" 2>/dev/null | column -t; } — chỉ nhằm ẩn log hệ thống và căn cột cho dễ đọc, không làm thay đổi kết quả truy vấn.')
add_image('sales_report.png', 'Bảng sales_report (Hive) — fact đa nguồn đã chuẩn hóa (POS+ERP+Ecom)')

heading('Các bảng Hive khác', 3)
table(
    ['Bảng', 'Nội dung', 'Cách ghi'],
    [
        ['dim_sanpham', '9 sản phẩm + unit_cost', 'Overwrite (dim nhỏ, cố định)'],
        ['inventory', 'Tồn kho = SUM(qty) nhập xuất', 'Overwrite (trạng thái hiện tại)'],
        ['snapshot_tonkho', 'Ảnh chụp tồn kho mỗi ngày', 'APPEND (dynamic partition)'],
        ['dim_khachhang', 'Khách hàng + phân khúc', 'Overwrite'],
    ]
)

# 2.4
heading('2.4. Tổng hợp Cascading (Incremental Rollup)', 2)

para('Một trong những thách thức lớn nhất khi xử lý dữ liệu theo chuỗi thời gian là đảm bảo mỗi khoảng thời gian chỉ được tổng hợp một lần duy nhất. Nếu không có cơ chế kiểm soát, việc chạy lại batch nhiều lần sẽ dẫn đến trùng lặp số liệu, làm sai lệch toàn bộ kết quả phân tích. Để giải quyết vấn đề này, hệ thống áp dụng nguyên lý cascading rollup với cơ chế watermark.')

heading('Nguyên lý', 3)
para('Mỗi kỳ — dù là ngày, tuần, tháng, quý hay năm — chỉ được tính đúng một lần, tại thời điểm kỳ đó đã đóng, nghĩa là không còn dữ liệu mới được thêm vào. Kỳ lớn được cuộn lên từ kỳ nhỏ hơn, không bao giờ đọc lại dữ liệu thô ban đầu. Điều này vừa đảm bảo tính chính xác, vừa giúp hệ thống hoạt động hiệu quả ngay cả khi dữ liệu tăng trưởng theo thời gian.')

heading('Sơ đồ Cascade', 3)
ascii_art("""sales_report (198,840 giao dịch ERP, ~398k dòng tổng)
     │
     ├──→ agg_ngay  (253 kỳ)    GROUP BY txn_date
     │       │
     │       ├──→ agg_tuan  (38 kỳ)    SUM(doanh_thu) GROUP BY week(ky)
     │       └──→ agg_thang (9 kỳ)     SUM(doanh_thu) GROUP BY month(ky)
     │                                    │
     │               ┌────────────────────┘
     │               ├──→ agg_quy  (3 kỳ)     SUM(doanh_thu) GROUP BY quarter(ky)
     │               └──→ agg_nam  (1 kỳ)     SUM(doanh_thu) GROUP BY year(ky)""")

para('Minh họa: agg_tuan của một tuần cụ thể được tính bằng tổng doanh thu của 7 ngày trong tuần đó từ bảng agg_ngay — không hề đọc lại bảng sales_report. Tương tự, agg_nam của năm 2026 được tính bằng tổng của 12 dòng tháng từ bảng agg_thang — không hề duyệt qua hàng trăm nghìn giao dịch thô.')

heading('Cơ chế Watermark', 3)
code('''def _append_closed_new(agg, ky_cur, table, src):
    agg = agg.where(F.col("ky") < ky_cur)      # chỉ kỳ đã đóng
    if has_table(table):
        wm = spark.sql(f"SELECT MAX(ky) FROM {table}").head()["k"]
        if wm is not None:
            agg = agg.where(F.col("ky") > wm)   # chỉ kỳ CHƯA xử lý
    agg.write.mode("append").saveAsTable(table)  # APPEND, không overwrite''')

para('Cơ chế watermark hoạt động như sau: trước khi thêm dữ liệu mới vào bảng agg_*, hệ thống kiểm tra giá trị ky lớn nhất đã có trong bảng đó. Chỉ những kỳ có ky lớn hơn watermark này mới được thêm vào. Nhờ vậy, dù batch có chạy lại 100 lần liên tiếp, kết quả vẫn không thay đổi — các kỳ cũ đã bị watermark chặn lại.')

heading('Kết quả Cascade', 3)
table(
    ['Bảng', 'Số kỳ', 'Từ ngày', 'Đến ngày', 'Nguồn cuộn từ'],
    [
        ['agg_ngay', '**253**', '2026-07-01', '2027-04-04', 'raw sales_report'],
        ['agg_tuan', '**38**', '2026-06-29', '2027-03-29', 'agg_ngay'],
        ['agg_thang', '**9**', '2026-07-01', '2027-03-01', 'agg_ngay'],
        ['agg_quy', '**3**', '2026-07-01', '2027-01-01', 'agg_thang'],
        ['agg_nam', '**1**', '2026-01-01', '2026-01-01', 'agg_thang'],
    ]
)

add_image('agg_thang.png', 'agg_thang (Hive) — cascade rollup doanh thu/lợi nhuận theo tháng')

# 2.5
heading('2.5. Đồng bộ kết quả ra PostgreSQL', 2)

para('Sau khi tất cả các bảng phân tích đã được tạo trong Hive, bước cuối cùng của batch layer là đẩy chúng sang PostgreSQL — nơi Grafana kết nối để trực quan hóa. Script spark_marts_to_pg.py thực hiện công việc này, sao chép 22 bảng từ Hive sang PostgreSQL.')

table(
    ['Loại bảng', 'Số bảng', 'Ví dụ'],
    [
        ['Báo cáo (bc_*)', '16', 'bc_doanhthu_cuahang, bc_loinhuan_vung, bc_top_sanpham...'],
        ['Tổng hợp (agg_*)', '5', 'agg_ngay, agg_tuan, agg_thang, agg_quy, agg_nam'],
        ['Dimension', '1', 'dim_khachhang'],
    ]
)

para('Bước này chạy ở chế độ cục bộ (không phân tán qua YARN) vì PostgreSQL chỉ được cài đặt trên master node. Nếu chạy trên executor tại slave, kết nối JDBC đến localhost:5432 sẽ trỏ nhầm sang chính slave đó thay vì master, dẫn đến lỗi kết nối. Đây là một quyết định thiết kế có chủ đích, không phải hạn chế của hệ thống.')

add_image('marts_output.png', 'spark_marts_to_pg — các bảng đã đẩy sang PostgreSQL')

# 2.6
heading('2.6. Cơ chế định kỳ và điều phối', 2)

para('Toàn bộ pipeline batch được lập lịch chạy tự động mỗi 15 phút thông qua Cron — trình lập lịch có sẵn trong hệ điều hành Linux. Không một bước nào trong pipeline yêu cầu can thiệp thủ công. Cấu hình crontab cụ thể như sau:')

code("*/15 * * * * bash -lc 'cd ~/big-data-final-project && bash run_batch.sh' >> ~/batch_cron.log 2>&1")

para('Trong thực tế doanh nghiệp, các pipeline Big Data thường được lập lịch bằng Apache Airflow hoặc Oozie — những công cụ quản lý workflow chuyên nghiệp với khả năng định nghĩa DAG phức tạp, retry khi lỗi, và cảnh báo qua email.Slack. Tuy nhiên, với pipeline gồm 5 bước tuần tự không có rẽ nhánh và hạ tầng 12GB đã vận hành 7 service, Cron kết hợp Bash là lựa chọn phù hợp: đơn giản, không tiêu tốn thêm tài nguyên, và tương đương về mặt nguyên lý vận hành. Trên thực tế, nhiều hệ thống production vẫn sử dụng cron cho các tác vụ định kỳ đơn giản.')

heading('Cơ chế khóa chống chồng', 3)
para('Một vấn đề có thể phát sinh khi lập lịch định kỳ là hai lần batch chạy chồng lên nhau — lần trước chưa kết thúc thì lần sau đã bắt đầu. Điều này đặc biệt nguy hiểm trên hệ thống có RAM hạn chế như cụm 12GB hiện tại. Để ngăn chặn, script run_batch.sh sử dụng cơ chế flock — file lock của Linux:')

code('exec 9>/tmp/run_batch.lock\nflock -n 9 || { echo "Batch trước chưa xong -> bỏ lần này."; exit 0; }')

para('Nếu lần batch trước chạy lâu hơn 15 phút, lần cron tiếp theo sẽ phát hiện lock và tự động bỏ qua, tránh tình trạng hai Spark job cùng chiếm dụng RAM và gây treo hệ thống.')

heading('Xử lý lỗi từng bước', 3)
para('Không phải bước nào trong pipeline cũng có mức độ quan trọng như nhau. Script run_batch.sh phân loại từng bước và có hành vi xử lý lỗi phù hợp:')

table(
    ['Bước', 'Mức độ', 'Hành vi khi lỗi'],
    [
        ['[1] Nạp Hive (spark_to_hive)', '**Quan trọng**', 'Hủy toàn bộ batch — các bước sau phụ thuộc vào sales_report'],
        ['[2] Incremental (spark_incremental)', '**Quan trọng**', 'Ghi nhận lỗi nhưng tiếp tục — phân tích kỳ vẫn được ưu tiên'],
        ['[3] Báo cáo SQL (spark_report_hive)', 'Nỗ lực tối đa', 'Bỏ qua, không chặn các bước sau'],
        ['[4] MLlib (spark_analysis)', 'Nỗ lực tối đa', 'Bỏ qua, mô hình có thể chạy lại vào lần sau'],
        ['[5] Đẩy Postgres (spark_marts_to_pg)', '**Quan trọng**', 'Luôn chạy cuối cùng, đảm bảo Grafana có dữ liệu'],
    ]
)

para('Kết thúc mỗi lần chạy, script ghi kết quả vào file batch_cron.log với dòng XONG (hoặc XONG (có lỗi: tên bước) nếu có bước thất bại). Cơ chế này giúp người quản trị dễ dàng theo dõi trạng thái hệ thống qua thời gian mà không cần truy cập vào giao diện giám sát phức tạp.')

# ── PART 3 ──
heading('PHẦN 3 — PHÂN TÍCH & TRỰC QUAN HÓA', 1)

# 3.1
heading('3.1. Phân tích kinh doanh với Spark SQL', 2)

para('Sau khi dữ liệu đã được nạp vào Hive và tổng hợp qua cascade, spark_report_hive.py thực thi một loạt truy vấn SQL để tạo ra các báo cáo kinh doanh. Tất cả báo cáo đều được lưu thành bảng Hive (bc_*) và đồng bộ sang PostgreSQL để Grafana trực quan hóa.')

heading('3.1.1. Doanh thu, chi phí, lợi nhuận theo cửa hàng và tháng', 3)
para('Đây là báo cáo chi tiết nhất, cho phép so sánh hiệu quả kinh doanh giữa các cửa hàng qua từng tháng. Mỗi dòng trong báo cáo thể hiện doanh thu, chi phí và lợi nhuận của một cửa hàng trong một tháng cụ thể. Dữ liệu được lấy từ ERP và lọc theo kênh offline (có store_id).')

table(
    ['Cửa hàng', 'Tháng', 'Doanh thu', 'Chi phí', 'Lợi nhuận'],
    [
        ['New York', '1/2027', '501,353', '348,769', '152,584'],
        ['Illinois', '1/2027', '493,989', '339,906', '154,083'],
        ['Michigan', '1/2027', '481,946', '331,745', '150,201'],
        ['Ohio', '1/2027', '458,593', '318,323', '140,270'],
        ['Georgia', '1/2027', '449,822', '316,107', '133,715'],
        ['California', '1/2027', '439,725', '311,333', '128,392'],
        ['Arizona', '1/2027', '432,701', '297,646', '135,055'],
    ]
)

para('Từ bảng trên có thể nhận thấy doanh thu giữa các cửa hàng khá đồng đều, dao động trong khoảng 430–500 nghìn mỗi tháng. Điều này phản ánh đặc điểm của tập dữ liệu Superstore gốc, nơi các khu vực và cửa hàng được thiết kế với quy mô tương đương nhau.')

add_image('bc_doanhthu_cuahang.png', 'Doanh thu/chi phí/lợi nhuận theo cửa hàng-tháng')

heading('3.1.2. Lợi nhuận theo khu vực', 3)
table(
    ['Khu vực', 'Doanh thu', 'Lợi nhuận'],
    [
        ['East', '**20,328,846**', '6,195,324'],
        ['South', '20,082,709', '6,143,478'],
        ['Central', '19,806,832', '6,094,188'],
        ['West', '19,748,014', '6,009,643'],
    ]
)

para('Bốn khu vực có doanh thu và lợi nhuận rất đồng đều, với East dẫn đầu nhẹ, chiếm khoảng 25.4% tổng doanh thu. Sự cân bằng này cho thấy hoạt động kinh doanh được phân bố tốt trên toàn bộ lãnh thổ, không phụ thuộc quá nhiều vào một khu vực duy nhất.')

add_image('bc_loinhuan_vung.png', 'Lợi nhuận theo khu vực')

heading('3.1.3. Top sản phẩm bán chạy', 3)
table(
    ['Sản phẩm', 'Doanh thu', 'Số lượng bán'],
    [
        ['**Copier** (TEC-CO-10004722)', '28,536,501', '66,138'],
        ['Machine (TEC-MA-10001127)', '18,966,040', '65,842'],
        ['Chair (FUR-CH-10002024)', '14,365,463', '66,482'],
        ['Bookcase (FUR-BO-10001798)', '11,496,084', '66,301'],
        ['Accessory (TEC-AC-10002049)', '3,841,537', '66,771'],
    ]
)

para('Một quan sát thú vị từ bảng trên là số lượng bán của tất cả sản phẩm đều xấp xỉ 66,000 đơn vị — gần như bằng nhau. Sự khác biệt về doanh thu đến từ đơn giá: Copier có giá 300/đơn vị, trong khi các sản phẩm văn phòng phẩm như Binder hay Paper chỉ có giá 5-8/đơn vị. Điều này cho thấy doanh thu được quyết định chủ yếu bởi cơ cấu sản phẩm hơn là số lượng bán.')

add_image('bc_top_sanpham.png', 'Top sản phẩm theo doanh thu')

heading('3.1.4. Doanh thu theo kênh', 3)
table(
    ['Kênh', 'Số giao dịch', 'Tỷ trọng'],
    [
        ['**Offline**', '126,616', '~64%'],
        ['Online', '72,224', '~36%'],
    ]
)

para('Kênh offline vẫn chiếm ưu thế với gần 2/3 tổng số giao dịch, phù hợp với mô hình doanh nghiệp bán lẻ truyền thống đang trong quá trình mở rộng lên trực tuyến. Cần lưu ý rằng POS không có customer_id nên không thể phân tích hành vi khách hàng đối với kênh offline — một hạn chế phổ biến trong thực tế bán lẻ.')

heading('3.1.5. Các báo cáo bổ sung', 3)
para('Ngoài các báo cáo chính đã trình bày, hệ thống còn tạo ra nhiều báo cáo phân tích khác, mỗi báo cáo phục vụ một khía cạnh cụ thể của hoạt động kinh doanh:')

table(
    ['Bảng', 'Nội dung'],
    [
        ['bc_doanhthu_segment', 'Doanh thu theo phân khúc khách (Consumer/Corporate/Home Office). Chỉ từ Ecommerce — JOIN sales_report với dim_khachhang qua customer_id'],
        ['bc_giatri_ton', 'Giá trị tồn kho theo sản phẩm (stock_qty × unit_cost)'],
        ['bc_kinhdoanh_ngay', 'Chuỗi thời gian doanh thu, lợi nhuận, số giao dịch theo ngày'],
        ['bc_online_thietbi', 'Doanh thu online theo thiết bị (mobile/desktop/tablet)'],
        ['bc_tuongquan', 'Tương quan giữa chi phí quảng cáo và doanh thu'],
        ['bc_xuhuong_tonkho', 'Xu hướng biến động tồn kho theo thời gian'],
        ['bc_canhbao_tonkho', 'Cảnh báo sản phẩm có tồn kho dưới ngưỡng an toàn'],
    ]
)

# 3.2
heading('3.2. Machine Learning — Spark MLlib', 2)

para('Bên cạnh các phân tích SQL truyền thống, hệ thống còn tích hợp thư viện MLlib của Spark để thực hiện các tác vụ học máy trên dữ liệu kinh doanh. Ba mô hình được triển khai, mỗi mô hình giải quyết một bài toán cụ thể trong vận hành doanh nghiệp.')

heading('3.2.1. Dự báo doanh thu — Linear Regression', 3)
para('Mô hình hồi quy tuyến tính được huấn luyện trên chuỗi thời gian 253 ngày để dự báo doanh thu cho tháng tiếp theo. Các đặc trưng đầu vào bao gồm ngày trong tuần, tháng, quý và doanh thu các kỳ trước. Kết quả dự báo được lưu vào bảng bc_dubao và hiển thị trên Grafana dưới dạng chỉ số thống kê.')

para('**Kết quả: Dự báo doanh thu tháng tiếp theo = 4,953,210**')
add_image('forecast.png', 'Dự báo doanh thu tháng tới (Linear Regression)')

heading('3.2.2. Phân cụm cửa hàng — K-Means', 3)
para('Mô hình K-Means được áp dụng để phân nhóm 12 cửa hàng dựa trên đặc điểm kinh doanh. Các đặc trưng sử dụng bao gồm doanh thu trung bình và số lượng giao dịch. Kết quả cho thấy 12 cửa hàng được phân thành 3 cụm riêng biệt, trong đó Georgia tách thành một cụm độc lập — cho thấy cửa hàng này có đặc điểm kinh doanh khác biệt so với phần còn lại.')

para('Kết quả một lần chạy minh họa (KMeans phụ thuộc khởi tạo ngẫu nhiên nên cụm có thể đổi giữa các lần chạy): Cụm 0 gồm Pennsylvania, California, New York, Michigan, Ohio; Cụm 1 gồm Washington, Illinois, Virginia, Arizona, Florida, Texas; Cụm 2: Georgia (riêng biệt).')

add_image('phancum.png', 'Phân cụm cửa hàng (K-Means)')

heading('3.2.3. Kế hoạch nhập hàng', 3)
para('Dựa trên dự báo doanh thu, tồn kho hiện tại và ngưỡng đặt lại của từng sản phẩm, hệ thống tự động đề xuất số lượng cần nhập. Đây là ứng dụng thực tế của việc kết hợp dự báo (Machine Learning) với dữ liệu vận hành (tồn kho) để hỗ trợ ra quyết định kinh doanh.')

add_image('kehoach_nhap.png', 'Kế hoạch nhập hàng đề xuất (kết hợp dự báo + tồn kho)')

# 3.3
heading('3.3. Streaming Real-time', 2)

para('Song song với batch layer chạy định kỳ, hệ thống còn vận hành một streaming layer đọc dữ liệu trực tiếp từ Kafka để cập nhật dashboard theo thời gian thực. Spark Streaming xử lý dữ liệu theo micro-batch, mỗi batch chỉ mất vài giây, cho phép phát hiện và phản ứng với các sự kiện gần như ngay lập tức.')

ascii_art("Kafka sales-report-clean\n     │\n     ▼\nSpark Streaming (micro-batch)\n     │\n     ├──→ rt_thongke  ( PostgreSQL → Grafana real-time )\n     └──→ rt_canhbao  ( PostgreSQL → Grafana alerts )")

table(
    ['Bảng', 'Nội dung', 'Tần suất cập nhật'],
    [
        ['rt_thongke', 'Doanh thu, lợi nhuận, số giao dịch theo thời gian thực', 'Liên tục (vài giây)'],
        ['rt_canhbao', 'Phát hiện giao dịch bất thường (lợi nhuận âm, doanh thu đột biến)', 'Liên tục (vài giây)'],
    ]
)

para('Streaming layer và batch layer chia sẻ cùng một nguồn dữ liệu từ Kafka, đảm bảo tính nhất quán. Trong khi batch layer cần 10-15 phút để hoàn thành một chu kỳ phân tích đầy đủ, streaming layer có thể cập nhật dashboard chỉ trong vài giây — một sự bổ trợ hoàn hảo cho kiến trúc xử lý dữ liệu.')

quote('Lưu ý về mốc thời gian: streaming ghi rt_* theo GIỜ XỬ LÝ THỰC (real-time) để trả lời "đang xảy ra gì ngay bây giờ"; còn batch dùng NGÀY NGHIỆP VỤ (txn_date) để phân tích lịch sử. Vì feeder nén thời gian (3 phút thật = 1 ngày nghiệp vụ) nên hai trục thời gian không trùng — đúng bản chất speed layer vs batch layer; trong hệ thống thật, giờ xử lý và ngày giao dịch trùng nhau.')

add_image('rt_thongke.png', 'rt_thongke — thống kê doanh thu/lợi nhuận real-time (streaming)')
add_image('rt_canhbao.png', 'rt_canhbao — cảnh báo giao dịch bất thường real-time')

# 3.4
heading('3.4. Trực quan hóa trên Grafana', 2)

para('Toàn bộ kết quả phân tích — cả từ batch layer lẫn streaming layer — được trực quan hóa trên Grafana tại địa chỉ http://192.168.79.131:3000. Hai dashboard riêng biệt phục vụ hai mục đích khác nhau: giám sát thời gian thực và phân tích kinh doanh chuyên sâu.')

heading('Dashboard 1: sales-report (Streaming)', 3)
para('Dashboard này tập trung vào giám sát hoạt động theo thời gian thực, hiển thị 4 panel chính:')

table(
    ['Panel', 'Loại biểu đồ', 'Nguồn dữ liệu'],
    [
        ['Doanh thu và Lợi nhuận', 'Timeseries', 'rt_thongke'],
        ['Số giao dịch', 'Timeseries', 'rt_thongke'],
        ['Cảnh báo gần đây', 'Table', 'rt_canhbao'],
        ['Tổng số cảnh báo', 'Stat', 'rt_canhbao'],
    ]
)

add_image('grafana_streaming.png', 'Dashboard sales-report (streaming) — 4 panel real-time')

heading('Dashboard 2: sales-report-batch (Batch)', 3)
para('Dashboard thứ hai là trung tâm phân tích của toàn bộ hệ thống, với 13 panel được sắp xếp theo thứ tự ưu tiên từ tổng quan đến chi tiết. Các panel timeseries cho phép người dùng kéo timeline để xem dữ liệu theo bất kỳ khoảng thời gian nào — từ một ngày cụ thể đến toàn bộ 9 tháng dữ liệu.')

table(
    ['Panel', 'Loại', 'Nguồn'],
    [
        ['TỔNG NĂM', 'Stat', 'agg_nam'],
        ['Dự báo tháng tới', 'Stat', 'bc_dubao'],
        ['**Doanh thu & LN theo NGÀY**', '**Timeseries**', 'agg_ngay'],
        ['Doanh thu theo kênh', 'Pie', 'bc_doanhthu_kenh'],
        ['Phân khúc khách', 'Pie', 'bc_doanhthu_segment'],
        ['Lợi nhuận theo vùng', 'Bar', 'bc_loinhuan_vung'],
        ['Top sản phẩm', 'Bar', 'bc_top_sanpham'],
        ['Tồn kho', 'Bar', 'bc_giatri_ton'],
        ['**Doanh thu & LN theo TUẦN**', '**Timeseries**', 'agg_tuan'],
        ['**Doanh thu & LN theo THÁNG**', '**Timeseries**', 'agg_thang'],
        ['**Doanh thu & LN theo QUÝ**', '**Timeseries**', 'agg_quy'],
        ['Kinh doanh NGÀY', 'Timeseries', 'bc_kinhdoanh_ngay'],
        ['Doanh thu theo tháng (bảng)', 'Table', 'bc_doanhthu_cuahang'],
    ]
)

para('Để xem dữ liệu theo thời gian, người dùng chọn time range "Last 2 years" trên thanh công cụ Grafana, sau đó kéo timeline đến khoảng thời gian mong muốn. Tất cả dữ liệu được cập nhật tự động sau mỗi 15 phút thông qua cron batch, hoặc theo thời gian thực qua streaming layer.')

add_image('grafana_batch.png', 'Dashboard sales-report-batch — 13 panel (Grafana)')

# ── PART 4 ──
heading('PHẦN 4 — TỔNG KẾT & KẾT LUẬN', 1)

# 4.1
heading('4.1. Tổng kết số liệu', 2)

heading('Dữ liệu', 3)
table(
    ['Chỉ số', 'Giá trị'],
    [
        ['Tổng giao dịch (từ ERP)', '**198,840**'],
        ['Tổng doanh thu', '79,966,401'],
        ['Tổng lợi nhuận', '24,442,633'],
        ['Phạm vi thời gian', '2026-07-01 → 2027-04-04 (253 ngày, ~9 tháng)'],
        ['Offline (POS)', '126,616 (~64%)'],
        ['Online (Ecom)', '72,224 (~36%)'],
        ['Số khách hàng định danh (CRM)', '53'],
        ['Số cửa hàng', '12'],
        ['Số khu vực', '4'],
        ['Số sản phẩm', '9'],
        ['Số ngày đã tổng hợp (agg_ngay)', '253'],
        ['Số tuần (agg_tuan)', '38'],
        ['Số tháng (agg_thang)', '9'],
        ['Số quý (agg_quy)', '3'],
        ['Số năm (agg_nam)', '1'],
    ]
)

heading('Dung lượng — Phân tầng Data Lake', 3)
table(
    ['Tầng', 'Mô tả', 'Định dạng', 'Dung lượng'],
    [
        ['**Bronze**', 'Raw JSON từ Kafka (/lake/transactions_archive/)', 'JSON Lines', '17,483 files, ~111 MB'],
        ['**Silver**', 'sales_report Hive table', 'Parquet (partition, nén)', 'Nhỏ hơn bronze'],
        ['**Gold**', 'agg_* + bc_* trong PostgreSQL', 'PostgreSQL tables', 'Vài MB'],
    ]
)

heading('Hiệu năng', 3)
table(
    ['Chỉ số', 'Giá trị'],
    [
        ['Tần suất batch', '15 phút/lần (cron)'],
        ['Thời gian 1 batch (YARN, 2 executor)', '~10-15 phút'],
        ['Dữ liệu mới mỗi batch (feeder tắt)', '0 (hệ thống ổn định)'],
        ['Streaming latency', 'Vài giây'],
        ['Tốc độ mô phỏng (khi feeder chạy)', '~1 ngày mô phỏng / 3 phút thật'],
    ]
)

# 4.2
heading('4.2. Kết luận và hướng phát triển', 2)

para('Qua quá trình xây dựng và vận hành hệ thống, nhóm đã đạt được những kết quả quan trọng sau:')

numbered('**Hệ thống xử lý dữ liệu hai lớp hoàn chỉnh** trên cụm Hadoop 3 node (6+3+3 GB), tích hợp đầy đủ Batch layer, Streaming layer và Serving layer. Dữ liệu được thu thập tự động từ 5 nguồn, xử lý qua NiFi và Kafka, lưu trữ phân tán trên HDFS, phân tích bằng Spark và trực quan hóa trên Grafana.')
numbered('**5 nguồn dữ liệu mô phỏng thực tế** doanh nghiệp đa kênh: POS, ERP, E-commerce, Kho và CRM. Mỗi nguồn được giả lập đúng cách hệ thống thật phát sinh dữ liệu, từ cơ chế real-time (POS), tích lũy database (ERP, Ecom, Kho) đến API (CRM).')
numbered('**Cơ chế Cascading Rollup với watermark** — điểm sáng về mặt kỹ thuật của hệ thống. Mỗi kỳ thời gian chỉ được tổng hợp một lần duy nhất khi đã đóng; các kỳ lớn được cuộn từ kỳ nhỏ. Cơ chế này đảm bảo tính chính xác, idempotent, và hiệu năng ổn định bất kể dữ liệu tăng trưởng bao lâu.')
numbered('**Incremental Ingest với cơ chế archive tự động** — mỗi lần batch chỉ đọc file mới, file cũ được chuyển sang archive. Pipeline không bị chậm dần theo thời gian, giữ được hiệu năng ổn định ngay cả khi tổng lượng dữ liệu đã lên đến hàng chục nghìn file.')
numbered('**Phân tầng Data Lake rõ ràng**: Bronze (JSON thô từ Kafka), Silver (Hive Parquet có cấu trúc, phân vùng), Gold (báo cáo tổng hợp trong PostgreSQL). Mỗi tầng phục vụ một mục đích khác nhau, từ lưu trữ thô đến phân tích.')
numbered('**Xử lý thành công 198,840 giao dịch ERP** (≈398k dòng đa nguồn) trên cụm RAM hạn chế (12GB tổng, 3 node). Con số này chứng minh hệ thống có thể vận hành hiệu quả ngay cả trong điều kiện tài nguyên eo hẹp.')
numbered('**Tự động hóa toàn bộ pipeline**: Cron lập lịch mỗi 15 phút, NiFi flow chạy liên tục, Spark Streaming xử lý real-time. Không một bước nào yêu cầu can thiệp thủ công sau khi khởi động.')
numbered('**Trực quan hóa đa chiều** với 17 panel Grafana trải dài từ tổng quan (năm, quý) đến chi tiết (ngày, sản phẩm). Người dùng có thể tự do cuộn timeline để xem dữ liệu ở bất kỳ cấp độ thời gian nào.')

para('Với hạ tầng hạn chế (12GB RAM, 3 node), hệ thống đã xử lý thành công 198,840 giao dịch ERP, cung cấp phân tích đa chiều theo ngày, tuần, tháng, quý và năm. Quan trọng hơn, kiến trúc này có thể mở rộng trực tiếp lên quy mô lớn hơn mà không cần thay đổi về mặt thiết kế — chỉ cần bổ sung tài nguyên phần cứng.')

heading('Hạn chế và hướng phát triển', 3)

table(
    ['Hạn chế', 'Nguyên nhân', 'Giải pháp'],
    [
        ['Phần cứng hạn chế (tổng 12GB, 3 node)', 'Môi trường bài tập', 'Trong phạm vi báo cáo, hệ thống đã xử lý thành công 198,840 giao dịch ERP. Có thể mở rộng lên quy mô lớn hơn khi tăng tài nguyên'],
        ['Dữ liệu mô phỏng từ Superstore', 'Không có dữ liệu doanh nghiệp thật', 'Đã làm giàu và phát sinh thành 5 nguồn riêng biệt, mô phỏng đúng đặc điểm từng hệ thống'],
        ['Spark khởi tạo cần tải thư viện mỗi lần (~200MB)', 'Chưa cấu hình spark.yarn.archive', 'Production có thể tải lên HDFS một lần, dùng lại cho mọi lần chạy'],
    ]
)

heading('Liên kết tham khảo', 3)
table(
    ['Thành phần', 'URL'],
    [
        ['Grafana', 'http://192.168.79.131:3000 (admin/admin)'],
        ['HDFS NameNode', 'http://192.168.79.131:9870'],
        ['YARN ResourceManager', 'http://192.168.79.131:8088'],
        ['NiFi', 'https://192.168.79.131:8443/nifi'],
        ['Source code', 'https://github.com/MacThiQuynhNhu/big-data-final-project'],
    ]
)

doc.add_paragraph()
quote('Báo cáo được tổng hợp từ dữ liệu thật đang chạy trên cụm Hadoop 3 node (master 10.0.2.195 + slave01 10.0.2.196 + slave02 10.0.2.197). Ngày sinh báo cáo: 06/2026. Dữ liệu phân tích cập nhật đến: 2027-04-04 (mô phỏng).')

# ── Save ──
add_page_numbers()
doc.save(DST)
print(f'Done: {DST}')
